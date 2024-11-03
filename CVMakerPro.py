import json
from docx import Document
from fpdf import FPDF

##CVMaker Pro, By Muhammad Usman Dilshad.

class PersonalInfo:
    def __init__(self, first_name, last_name, email, phone, country, city, nationality, job_title, birth_date, address,
                 postal_code, gender, hobbies, summary):
        self.first_name = first_name
        self.last_name = last_name
        self.email = email
        self.phone = phone
        self.country = country
        self.city = city
        self.nationality = nationality
        self.job_title = job_title
        self.birth_date = birth_date
        self.address = address
        self.postal_code = postal_code
        self.gender = gender
        self.hobbies = hobbies
        self.summary = summary

    def display(self):
        return f"Name: {self.first_name} {self.last_name}\nEmail: {self.email}\nContact: {self.phone}\nCountry: {self.country}\nCity: {self.city}\nNationality: {self.nationality}\nJob Title: {self.job_title}\nBirth Date: {self.birth_date}\nAddress: {self.address}\nPostal Code: {self.postal_code}\nGender: {self.gender}\nHobbies: {self.hobbies}\nProfessional Summary: {self.summary}"


class EmploymentHistory:
    def __init__(self):
        self.jobs = []

    def add_job(self, title, employer, designation, start_end_date, country, city, responsibilities):
        job = {
            "title": title,
            "employer": employer,
            "designation": designation,
            "start_end_date": start_end_date,
            "country": country,
            "city": city,
            "responsibilities": responsibilities
        }
        self.jobs.append(job)

    def display(self):
        result = ""
        for job in self.jobs:
            result += f"Job Title: {job['title']}\nEmployer: {job['employer']}\nDesignation: {job['designation']}\nDuration: {job['start_end_date']}\nCountry: {job['country']}\nCity: {job['city']}\nResponsibilities: {job['responsibilities']}\n\n"
        return result


class EducationDetail:
    def __init__(self):
        self.education = []

    def add_education(self, institution, degree, start_end_date, country, city, thesis):
        edu = {
            "institution": institution,
            "degree": degree,
            "start_end_date": start_end_date,
            "country": country,
            "city": city,
            "thesis": thesis
        }
        self.education.append(edu)

    def display(self):
        result = ""
        for edu in self.education:
            result += f"Institution: {edu['institution']}\nDegree: {edu['degree']}\nDuration: {edu['start_end_date']}\nCountry: {edu['country']}\nCity: {edu['city']}\nThesis: {edu['thesis']}\n\n"
        return result


class Skills:
    def __init__(self):
        self.skills = []

    def add_skill(self, skill_name, level):
        skill = {
            "skill_name": skill_name,
            "level": level
        }
        self.skills.append(skill)

    def display(self):
        result = ""
        for skill in self.skills:
            result += f"Skill: {skill['skill_name']}, Level: {skill['level']}\n"
        return result


class References:
    def __init__(self):
        self.references = []

    def add_reference(self, name, company, designation, email, phone):
        reference = {
            "name": name,
            "company": company,
            "designation": designation,
            "email": email,
            "phone": phone
        }
        self.references.append(reference)

    def display(self):
        result = ""
        for ref in self.references:
            result += f"Name: {ref['name']}, Company: {ref['company']}, Designation: {ref['designation']}, Email: {ref['email']}, Phone: {ref['phone']}\n"
        return result


class CVGenerator:
    def __init__(self):
        self.personal_info = None
        self.employment_history = EmploymentHistory()
        self.education_detail = EducationDetail()
        self.skills = Skills()
        self.references = References()

    def gather_personal_info(self):
        first_name = input("Enter your first name: ")
        last_name = input("Enter your last name: ")
        email = input("Enter your email: ")
        phone = input("Enter your phone number: ")
        country = input("Enter your country: ")
        city = input("Enter your city: ")
        nationality = input("Enter your nationality: ")
        job_title = input("Enter the job title you are applying for: ")
        birth_date = input("Enter your birth date: ")
        address = input("Enter your address: ")
        postal_code = input("Enter your postal code: ")
        gender = input("Enter your gender (M/F): ")
        hobbies = input("Enter your hobbies: ")
        summary = input("Enter a short professional summary: ")
        self.personal_info = PersonalInfo(first_name, last_name, email, phone, country, city, nationality, job_title,
                                          birth_date, address, postal_code, gender, hobbies, summary)

    def gather_employment_history(self):
        while True:
            title = input("Enter job title (type 'done' to stop): ")
            if title.lower() == 'done':
                break
            employer = input("Enter employer: ")
            designation = input("Enter designation: ")
            start_end_date = input("Enter start and end date: ")
            country = input("Enter country: ")
            city = input("Enter city: ")
            responsibilities = input("Enter responsibilities: ")
            self.employment_history.add_job(title, employer, designation, start_end_date, country, city,
                                            responsibilities)

    def gather_education_detail(self):
        while True:
            institution = input("Enter institution name (type 'done' to stop): ")
            if institution.lower() == 'done':
                break
            degree = input("Enter degree title: ")
            start_end_date = input("Enter start and end date: ")
            country = input("Enter country: ")
            city = input("Enter city: ")
            thesis = input("Enter thesis description: ")
            self.education_detail.add_education(institution, degree, start_end_date, country, city, thesis)

    def gather_skills(self):
        for _ in range(3):
            skill_name = input("Enter skill name: ")
            level = input("Enter skill level (low/intermediate/high): ")
            self.skills.add_skill(skill_name, level)

    def gather_references(self):
        add_references = input("Do you want to add references? (yes/no): ")
        if add_references.lower() == 'no':
            return
        for _ in range(3):
            name = input("Enter reference name: ")
            company = input("Enter company name: ")
            designation = input("Enter designation: ")
            email = input("Enter email: ")
            phone = input("Enter phone number: ")
            self.references.add_reference(name, company, designation, email, phone)

    def display_cv(self):
        print("\n=== CV ===\n")
        if self.personal_info:
            print("\n--- Personal Information ---\n")
            print(self.personal_info.display())
        print("\n--- Employment History ---\n")
        print(self.employment_history.display())
        print("\n--- Education Details ---\n")
        print(self.education_detail.display())
        print("\n--- Skills ---\n")
        print(self.skills.display())
        print("\n--- References ---\n")
        print(self.references.display())

    def export_to_word(self):
        doc = Document()
        doc.add_heading('Curriculum Vitae', 0)

        if self.personal_info:
            doc.add_heading('Personal Information', level=1)
            doc.add_paragraph(self.personal_info.display())

        doc.add_heading('Employment History', level=1)
        doc.add_paragraph(self.employment_history.display())

        doc.add_heading('Education Details', level=1)
        doc.add_paragraph(self.education_detail.display())

        doc.add_heading('Skills', level=1)
        doc.add_paragraph(self.skills.display())

        doc.add_heading('References', level=1)
        doc.add_paragraph(self.references.display())

        doc.save('CV_Output.docx')

    def export_to_pdf(self):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Curriculum Vitae", ln=True, align='C')

        if self.personal_info:
            pdf.cell(200, 10, txt="Personal Information", ln=True, align='L')
            pdf.multi_cell(0, 10, txt=self.personal_info.display())

        pdf.cell(200, 10, txt="Employment History", ln=True, align='L')
        pdf.multi_cell(0, 10, txt=self.employment_history.display())

        pdf.cell(200, 10, txt="Education Details", ln=True, align='L')
        pdf.multi_cell(0, 10, txt=self.education_detail.display())

        pdf.cell(200, 10, txt="Skills", ln=True, align='L')
        pdf.multi_cell(0, 10, txt=self.skills.display())

        pdf.cell(200, 10, txt="References", ln=True, align='L')
        pdf.multi_cell(0, 10, txt=self.references.display())

        pdf.output("CV_Output.pdf")

    def save_cv_versions(self, filename="saved_cvs.json"):
        cv_data = {
            "personal_info": vars(self.personal_info) if self.personal_info else {},
            "employment_history": self.employment_history.jobs,
            "education_detail": self.education_detail.education,
            "skills": self.skills.skills,
            "references": self.references.references
        }
        with open(filename, "w") as f:
            json.dump(cv_data, f)

    def load_cv_versions(self, filename="saved_cvs.json"):
        with open(filename, "r") as f:
            cv_data = json.load(f)
            if "personal_info" in cv_data:
                pi = cv_data["personal_info"]
                self.personal_info = PersonalInfo(**pi)
            for job in cv_data["employment_history"]:
                self.employment_history.add_job(**job)
            for edu in cv_data["education_detail"]:
                self.education_detail.add_education(**edu)
            for skill in cv_data["skills"]:
                self.skills.add_skill(**skill)
            for ref in cv_data["references"]:
                self.references.add_reference(**ref)

    def run(self):
        while True:
            print("\n1. Create a new CV")
            print("2. View CV")
            print("3. Export CV to Word")
            print("4. Export CV to PDF")
            print("5. Save CV")
            print("6. Load CV")
            print("7. Exit")
            choice = input("Enter your choice: ")

            if choice == '1':
                self.gather_personal_info()
                self.gather_employment_history()
                self.gather_education_detail()
                self.gather_skills()
                self.gather_references()
            elif choice == '2':
                self.display_cv()
            elif choice == '3':
                self.export_to_word()
                print("CV exported to Word format.")
            elif choice == '4':
                self.export_to_pdf()
                print("CV exported to PDF format.")
            elif choice == '5':
                self.save_cv_versions()
                print("CV saved successfully.")
            elif choice == '6':
                self.load_cv_versions()
                print("CV loaded successfully.")
            elif choice == '7':
                break
            else:
                print("Invalid choice. Please try again.")


if __name__ == "__main__":
    cv_generator = CVGenerator()
    cv_generator.run()
